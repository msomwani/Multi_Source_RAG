from fastapi import APIRouter, UploadFile, HTTPException, File, Query
from io import BytesIO
from typing import List, Dict
import json

import requests
from bs4 import BeautifulSoup
from PyPDF2 import PdfReader
import docx
import pdfplumber

from PIL import Image, ImageEnhance, ImageFilter
import pytesseract
from pdf2image import convert_from_bytes

from app.ingestion.text_splitter import chunk_text
from app.ingestion.table_utils import make_table_json, table_to_row_chunks
from app.llm.embeddings import embed
from app.vectorstore.store import add_chunks

router = APIRouter()

# --------------------------------------------------
# Loaders
# --------------------------------------------------
def load_pdf_bytes(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))
    return "\n".join(p.extract_text() or "" for p in reader.pages)


def load_docx_bytes(data: bytes) -> str:
    document = docx.Document(BytesIO(data))
    return "\n".join(p.text for p in document.paragraphs)


def load_docx_tables_structured(data: bytes) -> List[Dict]:
    document = docx.Document(BytesIO(data))
    tables = []

    for t_index, table in enumerate(document.tables):
        rows = [
            [(cell.text or "").strip().replace("\n", " ") for cell in row.cells]
            for row in table.rows
        ]

        if len(rows) < 2:
            continue

        tables.append(
            {
                "title": f"DOCX TABLE {t_index + 1}",
                "rows": rows,
            }
        )

    return tables


def load_pdf_tables_structured(
    data: bytes, max_pages: int = 10
) -> List[Dict]:
    tables_out = []

    with pdfplumber.open(BytesIO(data)) as pdf:
        for p_index, page in enumerate(pdf.pages[:max_pages]):
            for t_index, table in enumerate(page.extract_tables() or []):
                if not table or len(table) < 2:
                    continue

                rows = [
                    [(c or "").strip().replace("\n", " ") for c in row]
                    for row in table
                ]

                tables_out.append(
                    {
                        "title": f"PDF TABLE (Page {p_index + 1}, Table {t_index + 1})",
                        "rows": rows,
                    }
                )

    return tables_out


def load_web_page(url: str) -> str:
    response = requests.get(url, timeout=15)
    response.raise_for_status()

    soup = BeautifulSoup(response.text, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.extract()

    text = soup.get_text(separator="\n")
    return "\n".join(line.strip() for line in text.splitlines() if line.strip())


# --------------------------------------------------
# OCR helpers
# --------------------------------------------------
def preprocess_for_ocr(img: Image.Image) -> Image.Image:
    img = img.convert("L")
    img = ImageEnhance.Contrast(img).enhance(2.0)
    img = img.filter(ImageFilter.SHARPEN)

    w, h = img.size
    img = img.resize((w * 2, h * 2))

    return img


def load_image_bytes_ocr(data: bytes) -> str:
    img = Image.open(BytesIO(data)).convert("RGB")
    img = preprocess_for_ocr(img)
    return pytesseract.image_to_string(
        img, config="--oem 3 --psm 6"
    ).strip()


def load_pdf_bytes_with_ocr_fallback(
    data: bytes, max_pages: int = 15
) -> str:
    extracted = load_pdf_bytes(data)
    if extracted.strip():
        return extracted

    images = convert_from_bytes(
        data, dpi=300, first_page=1, last_page=max_pages
    )

    texts = []
    for idx, img in enumerate(images):
        img = preprocess_for_ocr(img)
        page_text = pytesseract.image_to_string(
            img, config="--oem 3 --psm 6"
        ).strip()

        if page_text:
            texts.append(f"\n\n--- Page {idx + 1} ---\n{page_text}")

    return "\n".join(texts).strip()


# --------------------------------------------------
# FILE INGEST
# --------------------------------------------------
@router.post("/ingest")
async def ingest_file(
    conversation_id: int = Query(...),
    file: UploadFile = File(...),
):
    raw = await file.read()
    name = file.filename.lower()

    chunks: list[str] = []
    metas: list[dict] = []

    text_main = ""
    tables_structured: list[Dict] = []

    if name.endswith(".pdf"):
        text_main = load_pdf_bytes_with_ocr_fallback(raw)
        tables_structured = load_pdf_tables_structured(raw)

    elif name.endswith(".docx"):
        text_main = load_docx_bytes(raw)
        tables_structured = load_docx_tables_structured(raw)

    elif name.endswith(".txt"):
        text_main = raw.decode("utf-8", errors="ignore").strip()

    elif name.endswith((".png", ".jpg", ".jpeg")):
        text_main = load_image_bytes_ocr(raw)

    else:
        raise HTTPException(400, "Unsupported file type")

    if not text_main and not tables_structured:
        raise HTTPException(400, "No text extracted from file")

    if text_main:
        main_chunks = chunk_text(text_main)
        chunks.extend(main_chunks)
        metas.extend(
            [{"source": file.filename, "type": "text"}]
            * len(main_chunks)
        )

    for table in tables_structured:
        table_json = make_table_json(table["title"], table["rows"])
        if not table_json:
            continue

        for rc in table_to_row_chunks(table_json):
            chunks.append(rc)
            metas.append(
                {
                    "source": file.filename,
                    "type": "table_row",
                    "table": json.dumps(table_json),
                    "table_title": table_json.get("title", ""),
                }
            )

    if not chunks:
        raise HTTPException(400, "No chunks produced from file")

    vectors = embed(chunks)

    add_chunks(
        file.filename,
        chunks,
        vectors,
        metas,
        conversation_id=conversation_id,
    )

    return {
        "status": "ok",
        "chunks": len(chunks),
        "tables": len(tables_structured),
    }


# --------------------------------------------------
# URL INGEST
# --------------------------------------------------
@router.post("/ingest/url")
async def ingest_url(
    conversation_id: int = Query(...),
    url: str = Query(...),
):
    text = load_web_page(url)
    if not text:
        raise HTTPException(400, "No text extracted from URL")

    chunks = chunk_text(text)
    vectors = embed(chunks)

    add_chunks(
        url,
        chunks,
        vectors,
        [{"source": url, "type": "text"}] * len(chunks),
        conversation_id=conversation_id,
    )

    return {
        "status": "ok",
        "source": url,
        "chunks": len(chunks),
    }
